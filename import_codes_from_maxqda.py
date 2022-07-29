import os
import docx
import re
import zipfile
import csv
import pandas as pd
import numpy as np
from docx.shared import RGBColor
from pathlib import Path
from datetime import datetime
from nltk import sent_tokenize
from bs4 import BeautifulSoup
from preprocess import (
    clean_sentence,
    remove_stop_words)


def import_codes(sentence2vec_model, doc_path, retrieved_codes_doc, theme_code_table_path, regexp):
    print(f'extracting sentences...')
    start = datetime.now()
    cat_df = pd.read_csv(theme_code_table_path,  encoding='utf-8-sig').applymap(lambda x: x.lower() if type(x) == str else x)
    cat_df.columns = cat_df.columns.str.lower()
    themes_found = []

    with zipfile.ZipFile(retrieved_codes_doc, 'r') as archive:
        train_columns = ['file_name', 'comment_id', 'original_sentence', 
            'cleaned_sentence', 'sentence_embedding', 'codes', 'themes']
        themes_list = list(cat_df)
        train_columns.extend(themes_list)

        train_df = pd.DataFrame(columns=train_columns)

        sentences_batch = []
        next_is_code = False
        missing_codes = []

        for paragraph in docx.Document(retrieved_codes_doc).paragraphs:
            # paragraph run is text from transcript
            if len(paragraph.runs) == 1:
                text = paragraph.runs[0].text
                text = text.replace('"', "'")
                text = text.replace('’', "'")

                for sentence in sent_tokenize(text):
                    sentences_batch.append(sentence)
            # paragraph runs is text from MAXQDA
            else:
                for run in paragraph.runs:
                    if next_is_code:
                        code = run.text.lower()
                        if '>' in code:
                            code = re.search(r'>(.*)', code).group(1).strip()
                        find_code = (cat_df.values == code).any(axis=0)
                        try:
                            current_theme = cat_df.columns[np.where(find_code==True)[0]].item()
                            if current_theme not in themes_found:
                                themes_found.append(current_theme)
                            # append all previous sentences to train table
                            for sentence in sentences_batch:
                                # sentence already in table, need to modify row
                                if train_df['original_sentence'].eq(sentence).any():
                                    # get matching row
                                    matching_index = train_df.index[train_df['original_sentence'] == sentence].tolist()[0]
                                    matching_row = train_df.iloc[matching_index]
                                    # update row
                                    matching_row_codes = [matching_code.lower() for matching_code in matching_row['codes'].split('; ')]
                                    matching_row_themes = [matching_theme.lower() for matching_theme in matching_row['themes'].split('; ')]

                                    if code not in matching_row_codes:
                                        matching_row_codes.append(code)
                                    
                                    if current_theme not in matching_row_themes:
                                        matching_row_themes.append(current_theme)

                                    new_codes = '; '.join(sorted(matching_row_codes, key=str.casefold))
                                    new_themes = '; '.join(sorted(matching_row_themes, key=str.casefold))
                                    
                                    train_df.loc[matching_index, 'codes'] = new_codes                                    
                                    train_df.loc[matching_index, 'themes'] = new_themes
                                    train_df.loc[matching_index, current_theme] = 1
                                # sentence not already in table, need to write row
                                else:
                                    cleaned_sentence = remove_stop_words(clean_sentence(sentence, regexp))
                                    row = {
                                        'file_name': re.search(r'([^\/]+).$', doc_path).group(0),
                                        'comment_id': '0', 
                                        'original_sentence': sentence, 
                                        'cleaned_sentence': cleaned_sentence, 
                                        'sentence_embedding': sentence2vec_model.get_vector(cleaned_sentence),
                                        'codes': code,
                                        'themes': current_theme
                                    }

                                    themes_binary = {}

                                    for theme in themes_list:
                                        if theme == current_theme:
                                            themes_binary[theme] = 1
                                        else:
                                            themes_binary[theme] = 0

                                    row.update(themes_binary)

                                    train_df = train_df.append(row, ignore_index=True)
                        except ValueError:
                            missing_codes.append(code)

                        next_is_code = False
                        sentences_batch = []
                        break
                    # char is '●'
                    elif '●' in run.text.strip():
                        next_is_code = True
    
    print(f'done extracting in {datetime.now() - start}')
    print(f'{len(set(missing_codes))} missing codes ({len(missing_codes)} sentences) in themes table (some counters in the codes table will be 0)')

    train_df.to_csv(doc_path.replace('.docx', '_train.csv'), index=False)
    
    print(f'themes found = {themes_found}')

    return themes_found


def create_codes_csv_from_maxqda(doc_path, retrieved_codes_doc):
    theme_code_table_path = os.path.join(Path(doc_path).parent.absolute(), doc_path.replace('.docx', '_codes.csv'))

    if not os.path.isfile(theme_code_table_path):
        print(f'extracting codes from from {retrieved_codes_doc}...')
        all_codes = []

        for paragraph in docx.Document(retrieved_codes_doc).paragraphs:
            code_search = re.search(r'(?<=Code: ● ).*', paragraph.text)
            if code_search:
                code_line = re.sub(r'(Weight score: \d+)$', '', code_search.group(0)).strip()
                if '>' in code_line:
                    code_line = re.search(r'>(.*)', code_line).group(1).strip()
                all_codes.append(code_line)

        codes_df = pd.DataFrame({'Theme 1 (replace this)': sorted(list(set(all_codes))), 'Theme 2 (replace this)': np.nan, 'Theme 3 (replace this)': np.nan, '...': np.nan})
        codes_df.to_csv(theme_code_table_path, index=False)

        print(f'{doc_path.replace(".docx", "_codes.csv")} created in {Path(doc_path).parent.absolute()}')
    else:
        print(f'code table already exists in {theme_code_table_path}')

    return theme_code_table_path