import itertools
import os
import re
import zipfile
import csv
import docx
import pandas as pd
import numpy as np
from pathlib import Path
from datetime import datetime
from nltk import sent_tokenize
from bs4 import BeautifulSoup
from preprocess import clean_sentence, remove_stop_words


# WORD
def create_codes_csv_from_word(doc_path, delimiter):
    theme_code_table_path = os.path.join(Path(doc_path).parent.absolute(), doc_path.replace('.docx', '_codes.csv'))

    if not os.path.isfile(theme_code_table_path):
        print(f'extracting codes from {doc_path}...')
        all_codes = []

        with zipfile.ZipFile(doc_path, 'r') as archive:
            comments_xml = archive.read('word/comments.xml')
            comments_soup = BeautifulSoup(comments_xml, 'xml')

            for comment in comments_soup.find_all('w:comment'):
                codes = comment.find_all('w:t')
                codes = ''.join([x.text for x in codes])
                codes = codes.strip().rstrip().lower()

                if delimiter != '' and delimiter in codes:
                    for code in codes.split(delimiter):
                        code = code.strip()
                        if len(code) > 0:
                            all_codes.append(code)
                else:
                    if len(codes) > 0:
                        all_codes.append(codes)

        codes_df = pd.DataFrame({'Theme 1 (replace this)': sorted(list(set(all_codes))), 'Theme 2 (replace this)': np.nan, 'Theme 3 (replace this)': np.nan, '...': np.nan})
        codes_df.to_csv(theme_code_table_path, index=False, encoding='utf-8-sig', errors='replace')

        print(f'{doc_path.replace(".docx", "_codes.csv")} created in {Path(doc_path).parent.absolute()}')
    else:
        print(f'code table already exists in {theme_code_table_path}')

    return theme_code_table_path


# NVIVO
def create_codes_csv_from_nvivo(doc_path, codes_folder_path):
    theme_code_table_path = os.path.join(codes_folder_path, doc_path.replace('.docx', '_codes.csv'))

    if not os.path.isfile(theme_code_table_path):
        print(f'extracting codes from {codes_folder_path}...')
        codes = []

        for entry in os.scandir(codes_folder_path):
            if entry.path.endswith('.docx'):
                code = entry.name[:-5].lower()
                codes.append(code)

        codes_df = pd.DataFrame({'Theme 1 (replace this)': sorted(list(set(codes))), 'Theme 2 (replace this)': np.nan, 'Theme 3 (replace this)': np.nan, '...': np.nan})
        codes_df.to_csv(theme_code_table_path, index=False, encoding='utf-8-sig', errors='replace')

        print(f'{doc_path.replace(".docx", "_codes.csv")} created in {codes_folder_path}')
    else:
        print(f'code table already exists in {theme_code_table_path}')

    return theme_code_table_path


# MAXQDA
def create_codes_csv_from_maxqda(doc_path, retrieved_codes_doc):
    theme_code_table_path = doc_path.replace('.docx', '_codes.csv')

    if not os.path.isfile(theme_code_table_path):
        print(f'extracting codes from {retrieved_codes_doc}...')
        all_codes = []

        for paragraph in docx.Document(retrieved_codes_doc).paragraphs:
            code_search = re.search(r'(?<=Code: ● ).*', paragraph.text)
            if code_search:
                code_line = re.sub(r'(Weight score: \d+)$', '', code_search.group(0)).strip()
                if '>' in code_line:
                    code_line = re.search(r'>(.*)', code_line).group(1).strip()
                all_codes.append(code_line)

        codes_df = pd.DataFrame({'Theme 1 (replace this)': sorted(list(set(all_codes))), 'Theme 2 (replace this)': np.nan, 'Theme 3 (replace this)': np.nan, '...': np.nan})
        codes_df.to_csv(theme_code_table_path, index=False, encoding='utf-8-sig', errors='replace')

        print(f'{doc_path.replace(".docx", "_codes.csv")} created in {Path(doc_path).parent.absolute()}')
    else:
        print(f'code table already exists in {theme_code_table_path}')

    return theme_code_table_path
    

# DEDOOSE
def create_codes_csv_from_dedoose(doc_path, excerpts_txt_path):
    theme_code_table_path = doc_path.replace('.docx', '_codes.csv')

    if not os.path.isfile(theme_code_table_path):
        print(f'extracting codes from {excerpts_txt_path}...')
        all_codes = []

        with open(excerpts_txt_path, errors='ignore', encoding='utf-8') as f:
            for line in f.readlines():
                codes_search = re.search(r'Codes Applied:\s(.*)\(\d+ of \d+-\d+\)', line)
                if codes_search:
                    codes = codes_search.group(0)
                    codes = re.sub(r'^Codes Applied:\s+', '', codes)
                    codes = re.sub(r'\t', ' ', codes)
                    codes = re.split(r'\(\d+ of \d+-\d+\)\s+', codes)
                    for code in codes:
                        if len(code) > 0:
                            code = re.sub(r'\([^)]+\)$', '', code)
                            all_codes.append(code)
            f.close()
        
        codes_df = pd.DataFrame({'Theme 1 (replace this)': sorted(list(set(all_codes))), 'Theme 2 (replace this)': np.nan, 'Theme 3 (replace this)': np.nan, '...': np.nan})
        codes_df.to_csv(theme_code_table_path, index=False, encoding='utf-8-sig', errors='replace')

        print(f'{doc_path.replace(".docx", "_codes.csv")} created in {Path(doc_path).parent.absolute()}')
    else:
        print(f'code table already exists in {theme_code_table_path}')

    return theme_code_table_path


# WORD
def transverse(start, end, text):
    if start == end:
        return text
    for node in itertools.chain([start], start.next_siblings):
        if node == end:
            # proper end
            return text

        if node.find('w:tab'):
            text += '\t'
        node = node.find('w:t')
        if node:
            text += node.text

    # if we get here it means we did not reach the end
    # ..so go up one level
    paragraph_node = start.parent

    # go to the next paragraph
    paragraph_siblings = paragraph_node.next_siblings
    next_paragraph = next(paragraph_siblings)
    # also count number of children to avoid StopIteration error in empty paragraphs
    while next_paragraph.name != 'w:p' or sum(1 for _ in next_paragraph.children) == 0:
        if next_paragraph == end:
            return text
        else:
            next_paragraph = next(paragraph_siblings)
    # this is a paragraph
    next_paragraph_rows = next_paragraph.children
    first_row = next(next_paragraph_rows)

    return transverse(first_row, end, text + '\n')


# WORD
# doc_path and theme_code_table_path documents already copied in data folder
def import_codes_from_word(sentence_bert, doc_path, delimiter, theme_code_table_path, regexp):

    start = datetime.now()
    cat_df = pd.read_csv(theme_code_table_path, encoding='utf-8-sig', encoding_errors='replace').applymap(lambda x: x.lower() if type(x) == str else x)
    cat_df.columns = cat_df.columns.str.lower()

    with zipfile.ZipFile(doc_path, 'r') as archive:
        # write header
        header = ['file_name', 'comment_id', 'original_sentence', 'cleaned_sentence', 'sentence_embedding', 'codes', 'themes']
        themes_list = list(cat_df)
        header.extend(themes_list)

        out_filename = doc_path.replace('.docx', '_train.csv')

        writer = csv.writer(open(out_filename, 'w', newline='', encoding='utf-8'))
        writer.writerow(header)

        doc_xml = archive.read('word/document.xml')
        doc_soup = BeautifulSoup(doc_xml, 'lxml')
        # open('tmp.xml', 'w').write(doc_soup.prettify())
        comments_xml = archive.read('word/comments.xml')
        comments_soup = BeautifulSoup(comments_xml, 'xml')

        missing_codes = []

        #print 'comments:'#, comments_soup.find_all('w:comment')
        for comment in comments_soup.find_all('w:comment'):
            codes = comment.find_all('w:t')
            codes = ''.join([x.text for x in codes])
            codes = codes.strip().rstrip().lower()
            comment_id = comment['w:id']

            # find range in doc_soup based on the comment id just found in comments_soup
            range_start = doc_soup.find('w:commentrangestart', attrs={'w:id': comment_id})
            range_end = doc_soup.find('w:commentrangeend', attrs={'w:id': comment_id})

            # transverse that range in doc_soup to extract the text that has been commented
            text = transverse(range_start, range_end, '')
            text = text.replace('"', "'")
            text = text.replace("’", "'")
            text = text.replace("´", "'")
            text = text.replace("…", "...")
            text = text.replace("\\", "\\\\")

            sentence_to_cleaned_dict = {}
            # split text into sentences
            for sentence in sent_tokenize(text):
                cleaned_sentence = remove_stop_words(clean_sentence(sentence, regexp))
                sentence_to_cleaned_dict[sentence] = [cleaned_sentence, sentence_bert.get_embeddings([cleaned_sentence])]

            themes = []

            if delimiter != '' and delimiter in codes:
                for code in codes.split(delimiter):
                    code = code.strip()
                    find_code = (cat_df.values == code).any(axis=0)
                    try:
                        theme = cat_df.columns[np.where(find_code==True)[0]].item()
                        if theme not in themes:
                            themes.append(theme)
                    except ValueError:
                        missing_codes.append(code)
                themes = sorted(themes, key=str.casefold)
                themes = '; '.join(themes)
            else:
                find_code = (cat_df.values == codes).any(axis=0)
                try:
                    themes = cat_df.columns[np.where(find_code==True)[0]].item()
                except ValueError:
                    missing_codes.append(codes)

            if len(themes) > 0:
                for sentence, tuple in sentence_to_cleaned_dict.items():
                    themes_binary = []
                    for theme in themes_list:
                        if theme in themes:
                            themes_binary.append(1)
                        else:
                            themes_binary.append(0)
                    row = [re.search(r'([^\/]+).$', doc_path).group(0), comment_id, sentence, tuple[0], tuple[1], codes, themes]
                    row.extend(themes_binary)
                    writer.writerow(row)

        print(f'done extracting in {datetime.now() - start}')

        print(f'{len(set(missing_codes))} missing codes ({len(missing_codes)} sentences) in themes table (some counters in the codes table will be 0)')
        # print(set(missing_codes))
        # os.remove('tmp.xml')
        return themes_list


# NVIVO
# doc_path and theme_code_table_path documents already copied in data folder
def import_codes_from_nvivo(sentence_bert, doc_path, codes_folder_path, theme_code_table_path, regexp):
    print(f'extracting codes from {codes_folder_path}...')
    start = datetime.now()

    # if theme_code_table_path == '':
    #     # create codes.csv in data folder from codes documents
    #     import_themes(doc_path, codes_folder_path)
    #     theme_code_table_path = doc_path.replace('.docx', '_codes.csv')

    cat_df = pd.read_csv(theme_code_table_path, encoding='utf-8-sig', encoding_errors='replace').applymap(lambda x: x.lower() if type(x) == str else x)
    cat_df.columns = cat_df.columns.str.lower()

    header = ['file_name', 'comment_id', 'original_sentence', 'cleaned_sentence', 'sentence_embedding', 'codes', 'themes']

    themes_list = [theme_name.lower() for theme_name in list(cat_df)]
    header.extend(themes_list)

    train_df = pd.DataFrame(columns=header)

    themes_found = []

    missing_codes = []

    for entry in os.scandir(codes_folder_path):
        if entry.path.endswith('.docx'):
            code = entry.name[:-5].lower()
            find_code = (cat_df.values == code).any(axis=0)
            theme = None

            try:
                theme = cat_df.columns[np.where(find_code==True)[0]].item().lower()
                if theme not in themes_found:
                    themes_found.append(theme)

                with zipfile.ZipFile(entry.path, 'r') as archive:
                    doc_xml = archive.read('word/document.xml')
                    doc_soup = BeautifulSoup(doc_xml, 'lxml')

                    for paragraph in doc_soup.find_all('w:p'):
                        if (paragraph.find('w:shd') is None and paragraph.find('w:highlight') is None and paragraph.find('w:t')):
                            text = paragraph.find('w:t').get_text() 
                            text = text.replace('"', "'")
                            text = text.replace('’', "'")
                            text = text.replace("´", "'")
                            text = text.replace("…", "...")
                            text = text.replace("\\", "\\\\")
                            for sentence in sent_tokenize(text):
                                cleaned_sentence = remove_stop_words(clean_sentence(sentence, regexp))
                            
                                if train_df['original_sentence'].eq(sentence).any():
                                    matching_index = train_df.index[train_df['original_sentence'] == sentence].tolist()[0]

                                    matching_row = train_df.iloc[matching_index]
                                    
                                    matching_row_codes = [matching_code.lower() for matching_code in matching_row['codes'].split('; ')]
                                    matching_row_themes = [matching_theme.lower() for matching_theme in matching_row['themes'].split('; ')]

                                    if code not in matching_row_codes:
                                        matching_row_codes.append(code)
                                    
                                    if theme not in matching_row_themes:
                                        matching_row_themes.append(theme)

                                    new_codes = '; '.join(sorted(matching_row_codes, key=str.casefold))
                                    new_themes = '; '.join(sorted(matching_row_themes, key=str.casefold))
                                    
                                    train_df.loc[matching_index, 'codes'] = new_codes                                    
                                    train_df.loc[matching_index, 'themes'] = new_themes
                                    train_df.loc[matching_index, theme] = 1
                                    
                                else:
                                    row = {
                                        'file_name': entry.path.replace('\\', '/'),
                                        'comment_id': '0', 
                                        'original_sentence': sentence,
                                        'cleaned_sentence': cleaned_sentence,
                                        'sentence_embedding': sentence_bert.get_embeddings([cleaned_sentence]),
                                        'codes': code,
                                        'themes': theme
                                    }

                                    theme_row = {}

                                    for theme_name in themes_list:
                                        if theme_name == theme:
                                            theme_row[theme_name] = 1
                                        else:
                                            theme_row[theme_name] = 0

                                    row.update(theme_row)

                                    train_df = train_df.append(row, ignore_index=True)
            except ValueError:
                missing_codes.append(code)            

    print(f'done extracting in {datetime.now() - start}')

    print(f'{len(set(missing_codes))} missing codes ({len(missing_codes)} sentences) in themes table (some counters in the codes table will be 0)')
    # print(set(missing_codes))

    train_df.to_csv(doc_path.replace('.docx', '_train.csv'), index=False, encoding='utf-8-sig', errors='replace')

    return themes_found


# MAXQDA
# doc_path and theme_code_table_path documents already copied in data folder
def import_codes_from_maxqda(sentence_bert, doc_path, retrieved_codes_doc, theme_code_table_path, regexp):
    print(f'extracting sentences...')
    start = datetime.now()
    cat_df = pd.read_csv(theme_code_table_path, encoding='utf-8-sig', encoding_errors='replace').applymap(lambda x: x.lower() if type(x) == str else x)
    cat_df.columns = cat_df.columns.str.lower()
    themes_found = []

    train_columns = ['file_name', 'comment_id', 'original_sentence', 'cleaned_sentence', 'sentence_embedding', 'codes', 'themes']
    themes_list = list(cat_df)
    train_columns.extend(themes_list)

    train_df = pd.DataFrame(columns=train_columns)

    sentences_batch = []
    next_is_code = False
    missing_codes = []

    for paragraph in docx.Document(retrieved_codes_doc).paragraphs:
        # paragraph run is text from transcript
        if len(paragraph.runs) == 1:
            text = paragraph.runs[0].text
            text = text.replace('"', "'")
            text = text.replace('’', "'")
            text = text.replace("´", "'")
            text = text.replace("…", "...")
            text = text.replace("\\", "\\\\")
            for sentence in sent_tokenize(text):
                sentences_batch.append(sentence)
        # paragraph runs is text from MAXQDA
        else:
            for run in paragraph.runs:
                if next_is_code:
                    code = run.text.lower()
                    if '>' in code:
                        code = re.search(r'>(.*)', code).group(1).strip()
                    find_code = (cat_df.values == code).any(axis=0)
                    try:
                        current_theme = cat_df.columns[np.where(find_code==True)[0]].item()
                        if current_theme not in themes_found:
                            themes_found.append(current_theme)
                        # append all previous sentences to train table
                        for sentence in sentences_batch:
                            # sentence already in table, need to modify row
                            if train_df['original_sentence'].eq(sentence).any():
                                # get matching row
                                matching_index = train_df.index[train_df['original_sentence'] == sentence].tolist()[0]
                                matching_row = train_df.iloc[matching_index]
                                # update row
                                matching_row_codes = [matching_code.lower() for matching_code in matching_row['codes'].split('; ')]
                                matching_row_themes = [matching_theme.lower() for matching_theme in matching_row['themes'].split('; ')]

                                if code not in matching_row_codes:
                                    matching_row_codes.append(code)
                                
                                if current_theme not in matching_row_themes:
                                    matching_row_themes.append(current_theme)

                                new_codes = '; '.join(sorted(matching_row_codes, key=str.casefold))
                                new_themes = '; '.join(sorted(matching_row_themes, key=str.casefold))
                                
                                train_df.loc[matching_index, 'codes'] = new_codes                                    
                                train_df.loc[matching_index, 'themes'] = new_themes
                                train_df.loc[matching_index, current_theme] = 1
                            # sentence not already in table, need to write row
                            else:
                                cleaned_sentence = remove_stop_words(clean_sentence(sentence, regexp))
                                row = {
                                    'file_name': re.search(r'([^\/]+).$', doc_path).group(0),
                                    'comment_id': '0', 
                                    'original_sentence': sentence, 
                                    'cleaned_sentence': cleaned_sentence, 
                                    'sentence_embedding': sentence_bert.getembeddings([cleaned_sentence]),
                                    'codes': code,
                                    'themes': current_theme
                                }

                                themes_binary = {}

                                for theme in themes_list:
                                    if theme == current_theme:
                                        themes_binary[theme] = 1
                                    else:
                                        themes_binary[theme] = 0

                                row.update(themes_binary)

                                train_df = train_df.append(row, ignore_index=True)
                    except ValueError:
                        missing_codes.append(code)

                    next_is_code = False
                    sentences_batch = []
                    break
                # char is '●'
                elif '●' in run.text.strip():
                    next_is_code = True
    
    print(f'done extracting in {datetime.now() - start}')
    print(f'{len(set(missing_codes))} missing codes ({len(missing_codes)} sentences) in themes table')

    train_df.to_csv(doc_path.replace('.docx', '_train.csv'), index=False, encoding='utf-8-sig', errors='replace')
    
    # print(f'themes found = {themes_found}')

    return themes_found


# DEDOOSE
# doc_path and theme_code_table_path documents already copied in data folder
def import_codes_from_dedoose(sentence_bert, doc_path, excerpts_txt_path, theme_code_table_path, regexp):
    print(f'extracting sentences...')
    start = datetime.now()
    cat_df = pd.read_csv(theme_code_table_path, encoding='utf-8-sig', encoding_errors='replace').applymap(lambda x: x.lower() if type(x) == str else x)
    cat_df.columns = cat_df.columns.str.lower()
    themes_found = []
    missing_codes = []

    train_columns = ['file_name', 'comment_id', 'original_sentence', 'cleaned_sentence', 'sentence_embedding', 'codes', 'themes']
    themes_list = list(cat_df)
    train_columns.extend(themes_list)

    train_df = pd.DataFrame(columns=train_columns)

    with open(excerpts_txt_path, errors='ignore', encoding='utf-8') as f:
        lines = f.readlines()
        lines_length = len(lines)
        current_codes = []
        current_themes = []
        excerpt_found = False
        current_excerpt = ''

        for i, line in enumerate(lines):
            # looking for code names
            if not current_codes:
                codes_search = re.search(r'Codes Applied:\s(.*)\(\d+ of \d+-\d+\)', line)
                if codes_search:
                    codes = codes_search.group(0)
                    codes = re.sub(r'^Codes Applied:\s+', '', codes)
                    codes = re.sub(r'\t', ' ', codes)
                    codes = re.split(r'\(\d+ of \d+-\d+\)\s+', codes)
                    for code in codes:
                        if len(code) > 0:
                            code = re.sub(r'\([^)]+\)$', '', code)
                            code = code.lower()
                            current_codes.append(code)
                            find_code = (cat_df.values == code).any(axis=0)
                            try:
                                theme = cat_df.columns[np.where(find_code==True)[0]].item()
                                if theme not in current_themes:
                                    current_themes.append(theme)
                                if theme not in themes_found:
                                    themes_found.append(theme)
                            except ValueError:
                                missing_codes.append(code)
            # looking for excerpt start
            elif not excerpt_found:
                if re.search(r'^Excerpt Range: \d+-\d+', line):
                    excerpt_found = True
            # looking for excerpt text
            else:
                # continue reading if not at the end of the excerpt or file
                if i < lines_length-1 and (not re.search(r'^Title: ', line) or not re.search(r'^Doc Creator: ', lines[i+1])):
                    current_excerpt += line
                else:
                    current_excerpt = current_excerpt.replace('"', "'")
                    current_excerpt = current_excerpt.replace('’', "'")
                    current_excerpt = current_excerpt.replace("´", "'")
                    current_excerpt = current_excerpt.replace("…", "...")
                    current_excerpt = current_excerpt.replace("\\", "\\\\")

                    for sentence in sent_tokenize(current_excerpt):
                        cleaned_sentence = remove_stop_words(clean_sentence(sentence, regexp))
                        row = {
                            'file_name': re.search(r'([^\/]+).$', doc_path).group(0),
                            'comment_id': '0',
                            'original_sentence': sentence,
                            'cleaned_sentence': cleaned_sentence,
                            'sentence_embedding': sentence_bert.get_embeddings(cleaned_sentence),
                            'codes': '; '.join(current_codes),
                            'themes': '; '.join(current_themes)                            
                        }

                        themes_binary = {}

                        for theme in themes_list:
                            if theme in current_themes:
                                themes_binary[theme] = 1
                            else:
                                themes_binary[theme] = 0

                        row.update(themes_binary)

                        train_df = train_df.append(row, ignore_index=True)

                    current_codes = []
                    current_themes = []
                    excerpt_found = False
                    current_excerpt = ''  

        print(f'done extracting in {datetime.now() - start}')
        print(f'{len(set(missing_codes))} missing codes ({len(missing_codes)} sentences) in themes table')

        train_df.to_csv(doc_path.replace('.docx', '_train.csv'), index=False, encoding='utf-8-sig', errors='replace')

        # print(f'themes found = {themes_found}')

        f.close()

        return themes_found